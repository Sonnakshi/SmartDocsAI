import io
import re
import urllib.parse
import requests
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# DOCX formatting
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# PDF formatting
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors


# ==========================================
# 1. TEXT SANITIZATION & UNICODE CLEANER
# ==========================================

def sanitize_unicode(text: str) -> str:
    """Replaces Unicode special characters with standard clean characters."""
    if not text:
        return ""
    text = text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    text = text.replace("—", " - ").replace("–", " - ").replace("‑", "-").replace("−", "-")
    text = text.replace("•", "* ").replace("…", "...").replace("·", "-")
    text = text.replace("\u00a0", " ").replace("\u2009", " ").replace("\t", "    ")
    return text


def clean_markdown_to_xml(text: str) -> str:
    """Converts markdown bold, italics, code, and math into clean XML for PDF."""
    text = sanitize_unicode(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)
    text = re.sub(r"`(.*?)`", r'<font name="Courier" color="#0369a1">\1</font>', text)
    text = re.sub(r"\$(.*?)\$", r"<i>\1</i>", text)
    text = text.replace("<br>", "<br/>")
    return text


# ==========================================
# 2. AI IMAGE GENERATION
# ==========================================

def generate_ai_image(prompt: str) -> bytes:
    """Generates an image via real-time diffusion models and returns raw PNG/JPEG bytes."""
    clean_prompt = urllib.parse.quote(prompt.strip())
    url = f"https://image.pollinations.ai/prompt/{clean_prompt}?width=1024&height=576&nologo=true"
    try:
        res = requests.get(url, timeout=30)
        if res.status_code == 200:
            return res.content
    except Exception:
        pass
    return None


# ==========================================
# 3. DYNAMIC MATHEMATICAL GRAPH PLOTTING
# ==========================================

def extract_and_render_plots(text: str):
    """Detects Python matplotlib code blocks in the AI response and renders high-res PNG plots."""
    code_pattern = r"```python(.*?)```"
    matches = re.findall(code_pattern, text, re.DOTALL)
    figures_bytes = []
    
    for code in matches:
        if "plt." in code or "matplotlib" in code:
            try:
                plt.figure(figsize=(8, 4.5), dpi=150)
                plt.style.use("seaborn-v0_8-whitegrid" if "seaborn-v0_8-whitegrid" in plt.style.available else "default")
                local_scope = {"plt": plt}
                exec(code, {}, local_scope)
                
                buf = io.BytesIO()
                plt.savefig(buf, format="png", bbox_inches="tight")
                buf.seek(0)
                figures_bytes.append(buf.getvalue())
                plt.close("all")
            except Exception:
                plt.close("all")
                
    return figures_bytes


# ==========================================
# 4. PUBLICATION-GRADE PDF EXPORT
# ==========================================

def generate_pdf_export(question: str, answer: str, citations: list) -> bytes:
    """Generates an executive, publication-styled PDF report with tables, headings, and citations."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=48, leftMargin=48, topMargin=48, bottomMargin=48)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle("DocTitle", fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#0f172a"), spaceAfter=6)
    sub_style = ParagraphStyle("DocSub", fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#64748b"), spaceAfter=14)
    q_label = ParagraphStyle("QLabel", fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#0284c7"), spaceAfter=4)
    q_text = ParagraphStyle("QText", fontName="Helvetica-Oblique", fontSize=11, leading=15, textColor=colors.HexColor("#1e293b"), spaceAfter=12)
    h2_style = ParagraphStyle("Heading2", fontName="Helvetica-Bold", fontSize=13, leading=17, textColor=colors.HexColor("#0f172a"), spaceBefore=12, spaceAfter=6)
    h3_style = ParagraphStyle("Heading3", fontName="Helvetica-Bold", fontSize=11, leading=15, textColor=colors.HexColor("#334155"), spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontName="Helvetica", fontSize=9.5, leading=14, textColor=colors.HexColor("#334155"), spaceAfter=6)
    bullet_style = ParagraphStyle("Bullet", fontName="Helvetica", fontSize=9.5, leading=14, textColor=colors.HexColor("#334155"), leftIndent=12, spaceAfter=4)
    cite_head = ParagraphStyle("CiteHead", fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#0f172a"), spaceBefore=14, spaceAfter=6)
    cite_item = ParagraphStyle("CiteItem", fontName="Helvetica", fontSize=8.5, leading=12, textColor=colors.HexColor("#64748b"), spaceAfter=3)
    table_cell = ParagraphStyle("TableCell", fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#334155"))
    table_hdr = ParagraphStyle("TableHdr", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.HexColor("#ffffff"))

    story = [
        Paragraph("SmartDocs AI - Intelligence Report", title_style),
        Paragraph("Automated Document Synthesis & RAG Analysis", sub_style),
        HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#e2e8f0"), spaceAfter=12),
        Paragraph("User Query:", q_label),
        Paragraph(sanitize_unicode(question), q_text),
        HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=12),
    ]

    clean_answer = re.sub(r"```python(.*?)```", "", answer, flags=re.DOTALL)
    clean_answer = re.sub(r"\[IMAGE_PROMPT:.*?\]", "", clean_answer).strip()

    lines = clean_answer.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("### "):
            story.append(Paragraph(clean_markdown_to_xml(line[4:]), h3_style))
            i += 1
            continue

        if line.startswith("## ") or line.startswith("# "):
            clean_hdr = re.sub(r"^#+\s*", "", line)
            story.append(Paragraph(clean_markdown_to_xml(clean_hdr), h2_style))
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                t_line = lines[i].strip()
                if not re.match(r"^\|[\s\-:|]+\|$", t_line):
                    cells = [c.strip() for c in t_line.split("|")[1:-1]]
                    table_lines.append(cells)
                i += 1

            if table_lines:
                formatted_table_data = []
                for row_idx, row in enumerate(table_lines):
                    formatted_row = []
                    for cell in row:
                        style = table_hdr if row_idx == 0 else table_cell
                        formatted_row.append(Paragraph(clean_markdown_to_xml(cell), style))
                    formatted_table_data.append(formatted_row)

                col_count = len(table_lines[0])
                col_width = 516 / max(col_count, 1)
                t = Table(formatted_table_data, colWidths=[col_width] * col_count)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0284c7")),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#ffffff"), colors.HexColor("#f8fafc")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(Spacer(1, 4))
                story.append(t)
                story.append(Spacer(1, 8))
            continue

        if re.match(r"^(\*|-|•|\d+\.)\s+", line):
            bullet_text = re.sub(r"^(\*|-|•|\d+\.)\s+", "", line)
            story.append(Paragraph(f"&bull; {clean_markdown_to_xml(bullet_text)}", bullet_style))
            i += 1
            continue

        story.append(Paragraph(clean_markdown_to_xml(line), body_style))
        i += 1

    if citations:
        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=8))
        story.append(Paragraph("Verified Source Citations", cite_head))
        for idx, c in enumerate(citations, 1):
            fn = sanitize_unicode(c.get("filename", "Doc"))
            chk = c.get("chunk_index", 0)
            sc = c.get("score", 0)
            story.append(Paragraph(f"[{idx}] <b>{fn}</b> — Chunk #{chk} (Similarity Score: {sc:.4f})", cite_item))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# 5. ENTERPRISE WORD (.DOCX) EXPORT WITH REAL TABLES
# ==========================================

def set_cell_background(cell, fill_hex):
    """Sets background color of a Word table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def format_docx_paragraph(p, text):
    """Parses markdown bold, italics, code, and adds runs to a Word paragraph."""
    text = sanitize_unicode(text)
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = p.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = p.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(2, 132, 199)
        elif token.startswith("$") and token.endswith("$"):
            run = p.add_run(token[1:-1])
            run.italic = True
            run.font.color.rgb = RGBColor(15, 23, 42)
        else:
            p.add_run(token)


def generate_docx_export(question: str, answer: str, citations: list) -> bytes:
    """Generates an executive, publication-styled Word document (.docx) with real native tables."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("SmartDocs AI - Intelligence Report", level=0)
    title.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    title.runs[0].font.size = Pt(22)

    p_q = doc.add_paragraph()
    run_lbl = p_q.add_run("User Query: ")
    run_lbl.bold = True
    run_lbl.font.color.rgb = RGBColor(2, 132, 199)
    run_q = p_q.add_run(sanitize_unicode(question))
    run_q.italic = True
    run_q.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_heading("Synthesized Analysis", level=1)
    
    clean_answer = re.sub(r"```python(.*?)```", "", answer, flags=re.DOTALL)
    clean_answer = re.sub(r"\[IMAGE_PROMPT:.*?\]", "", clean_answer).strip()

    lines = clean_answer.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("### "):
            h = doc.add_heading(level=3)
            format_docx_paragraph(h, line[4:])
            i += 1
            continue

        if line.startswith("## ") or line.startswith("# "):
            clean_hdr = re.sub(r"^#+\s*", "", line)
            h = doc.add_heading(level=2)
            format_docx_paragraph(h, clean_hdr)
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                t_line = lines[i].strip()
                if not re.match(r"^\|[\s\-:|]+\|$", t_line):
                    cells = [c.strip() for c in t_line.split("|")[1:-1]]
                    table_lines.append(cells)
                i += 1

            if table_lines:
                rows = len(table_lines)
                cols = len(table_lines[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"

                for r_idx, row_data in enumerate(table_lines):
                    row = table.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_data):
                        cell = row.cells[c_idx]
                        p = cell.paragraphs[0]
                        format_docx_paragraph(p, cell_text)
                        
                        if r_idx == 0:
                            set_cell_background(cell, "0284C7")
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            if r_idx % 2 == 0:
                                set_cell_background(cell, "F8FAFC")
                doc.add_paragraph()
            continue

        if re.match(r"^(\*|-|•|\d+\.)\s+", line):
            bullet_text = re.sub(r"^(\*|-|•|\d+\.)\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            format_docx_paragraph(p, bullet_text)
            i += 1
            continue

        p = doc.add_paragraph()
        format_docx_paragraph(p, line)
        i += 1

    if citations:
        doc.add_heading("Verified Source Citations", level=1)
        for idx, c in enumerate(citations, 1):
            p = doc.add_paragraph()
            r1 = p.add_run(f"[{idx}] {sanitize_unicode(c.get('filename'))} ")
            r1.bold = True
            r1.font.color.rgb = RGBColor(2, 132, 199)
            p.add_run(f"(Chunk #{c.get('chunk_index')} | Similarity Score: {c.get('score', 0):.4f})\n")
            r2 = p.add_run(f"\"{sanitize_unicode(c.get('snippet', ''))}\"")
            r2.italic = True
            r2.font.color.rgb = RGBColor(100, 116, 139)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# 6. STRUCTURED TEXT (.TXT) EXPORT
# ==========================================

def generate_txt_export(question: str, answer: str, citations: list) -> str:
    """Generates an executive ASCII structured text report."""
    clean_answer = re.sub(r"```python(.*?)```", "", answer, flags=re.DOTALL)
    clean_answer = re.sub(r"\[IMAGE_PROMPT:.*?\]", "", clean_answer).strip()
    clean_answer = sanitize_unicode(clean_answer)

    divider = "=" * 76
    sub_divider = "-" * 76

    txt = [
        divider,
        "SMARTDOCS AI - INTELLIGENCE REPORT",
        "Document Synthesis & RAG Analysis",
        divider,
        f"USER QUERY:\n{sanitize_unicode(question)}",
        sub_divider,
        "SYNTHESIZED ANALYSIS:",
        clean_answer,
    ]

    if citations:
        txt.append(sub_divider)
        txt.append("VERIFIED SOURCE CITATIONS:")
        for idx, c in enumerate(citations, 1):
            txt.append(f"[{idx}] {sanitize_unicode(c.get('filename'))} (Chunk #{c.get('chunk_index')}, Score: {c.get('score', 0):.4f})")
            txt.append(f"    Excerpt: \"{sanitize_unicode(c.get('snippet', ''))}\"\n")

    txt.append(divider)
    return "\n\n".join(txt)